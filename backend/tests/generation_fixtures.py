"""Offline fixtures for the Phase 1 (form-fill) generation tests.

``make_fillable_pdf`` builds a small AcroForm PDF in-memory with reportlab so the
enumeration + source-upload paths can be exercised without any sample asset or network.
"""

from __future__ import annotations

import io

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


def make_fillable_pdf() -> bytes:
    """A one-page AcroForm PDF: two text fields, a checkbox, a dropdown, a signature.

    The last text field is literally named ``Signature`` so the name-based override to
    ``kind="signature"`` is covered.
    """
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    form = c.acroForm

    c.drawString(50, 700, "Vendor:")
    form.textfield(name="vendor_name", x=140, y=692, width=220, height=18)

    c.drawString(50, 660, "Total:")
    form.textfield(name="total_amount", x=140, y=652, width=220, height=18)

    c.drawString(50, 620, "Approved:")
    form.checkbox(name="approved", x=140, y=616, size=16)

    c.drawString(50, 580, "Currency:")
    form.choice(
        name="currency",
        value="USD",
        x=140,
        y=572,
        width=120,
        height=18,
        options=["USD", "EUR", "GBP"],
    )

    c.drawString(50, 540, "Signature:")
    form.textfield(name="Signature", x=140, y=532, width=220, height=18)

    c.save()
    return buf.getvalue()


def make_plain_pdf() -> bytes:
    """A one-page PDF with no AcroForm (drives the ``rich_html`` fallback path)."""
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    c.drawString(50, 700, "Just some prose, no form fields here.")
    c.save()
    return buf.getvalue()


# A distinctive heading string the docx conversion test asserts survives the round-trip.
DOCX_HEADING = "Purchase Agreement"


def make_docx_bytes() -> bytes:
    """A tiny in-memory .docx: a heading, a paragraph, and a 2-cell table.

    Built with python-docx (a docgen transitive dep via html-for-docx) so the rich-HTML
    conversion path can be exercised without any sample asset.
    """
    import docx  # python-docx

    document = docx.Document()
    document.add_heading(DOCX_HEADING, level=1)
    document.add_paragraph("This agreement is made between the parties named below.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Vendor"
    table.rows[0].cells[1].text = "Acme Supplies Inc."

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def make_xlsx_template() -> bytes:
    """A tiny in-memory .xlsx spreadsheet template for the spreadsheet-mode tests.

    Layout (sheet "Invoice"): a label/value scalar pair (``A1`` "Vendor" / ``B1`` blank);
    a line-items header (row 3) + one bold-styled anchor row (row 4) whose ``D4`` carries a
    per-row ``=B4*C4`` formula; and a ``SUM`` total below (``D6``). Exercises scalar fill,
    both table row modes (style clone + formula translation off the anchor row), and the
    LibreOffice recompute (the formulas have no cached result until recalc).
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font

    wb = Workbook()
    ws = wb.active
    ws.title = "Invoice"

    # Scalar label/value pair.
    ws["A1"] = "Vendor"
    ws["B1"] = None  # bound to `vendor`

    # Line-items header.
    ws["A3"] = "Description"
    ws["B3"] = "Qty"
    ws["C3"] = "Unit Price"
    ws["D3"] = "Amount"

    # Styled anchor row (row 4) with a per-row formula in the amount column.
    for letter in ("A", "B", "C", "D"):
        ws[f"{letter}4"].font = Font(bold=True)
    ws["C4"].number_format = '#,##0.00'
    ws["D4"].number_format = '#,##0.00'
    ws["D4"] = "=B4*C4"

    # Total below the anchor row (a formula with no cached result until recalc).
    ws["A6"] = "Total"
    ws["D6"] = "=SUM(D4:D4)"

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# A rich-HTML template body carrying the two placeholder marker kinds the binder fills:
# ``span[data-field]`` text placeholders (a scalar path + an indexed list path) and an
# ``img[data-signature]`` signature stamp target.
RICH_HTML_FIXTURE = (
    "<h1>Invoice</h1>"
    '<p>Vendor: <span data-field="vendor">Vendor</span></p>'
    '<p>Line amount: <span data-field="line_items.0.amount">Amt</span></p>'
    '<p>Missing: <span data-field="po_number">PO</span></p>'
    '<p>Signed: <img data-signature src="" alt="signature"></p>'
)
